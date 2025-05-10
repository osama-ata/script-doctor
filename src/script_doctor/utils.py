"""
utils.py
Utility functions for the script_doctor module.
"""

import os
from docx import Document


# a function to read all the .docx files in a directory
def read_docx_files_in_directory(directory):
    """
    Reads all .docx files in a given directory and returns their contents.

    Args:
        directory (str): The path to the directory containing .docx files.

    Returns:
        dict: A dictionary where keys are file names and values are lists of
        lists representing the tables.
    """

    # Check if the directory exists
    if not os.path.isdir(directory):
        raise ValueError("The provided path is not a valid directory")

    # Initialize an empty dictionary to store the contents of each file
    docx_contents = {}

    # Iterate over all files in the directory
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            file_path = os.path.join(directory, filename)
            docx_contents[filename] = read_docx_table(file_path)

    return docx_contents


# a function to read a table from a .docx file
def read_docx_table(file_obj_or_path):
    """
    Reads the first table from a .docx file (or file-like object)
    and returns it as a list of lists.

    Args:
        file_obj_or_path (str or file-like object): The path to the .docx file
                                                or a file-like object.

    Returns:
        list: A list of lists representing the table,
        or None if no table is found.
    """
    # Check if the file path ends with .docx if it's a string
    if isinstance(file_obj_or_path, str):
        if not file_obj_or_path.endswith(".docx"):
            raise ValueError("The file must be a .docx file")

    # Read the document and extract the first table
    try:
        # Document can take a path or a file-like object
        doc = Document(file_obj_or_path)
    except Exception as e:
        # Consider more specific error handling if possible
        raise IOError(f"Error reading the .docx file/object: {e}")
    # Check if the document contains any tables
    if not doc.tables:
        return None  # Return None if no tables are found

    # Read the first table
    table = doc.tables[0]
    return [[cell.text for cell in row.cells] for row in table.rows]
